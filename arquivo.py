from tkinter import filedialog, messagebox
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches
from util import corrigir_caracteres, converter_markdown
from PIL import Image

def salvar_como_docx(texto, imagem_path=None):
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Documentos do Word", "*.docx")])
    if not file_path:
        return

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    linhas = [corrigir_caracteres(l.strip()) for l in texto.split("\n")]
    if len(linhas) < 2:
        messagebox.showerror("Erro", "Texto inválido.")
        return

    nome, _ = converter_markdown(linhas[0])
    contato, _ = converter_markdown(linhas[1])

    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0]
    if imagem_path:
        try:
            paragraph = row.cells[0].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(imagem_path, width=Inches(1.2))
        except Exception as e:
            print("Erro ao inserir imagem no docx:", e)

    paragraph = row.cells[1].paragraphs[0]
    run = paragraph.add_run(nome)
    run.bold = True
    run.font.size = Pt(20)

    paragraph.add_run("\n" + contato)
    doc.add_paragraph("-" * 30)

    for linha in linhas[2:]:
        if not linha.strip():
            doc.add_paragraph()
            continue
        linha_formatada, is_negrito = converter_markdown(linha)
        if is_negrito:
            doc.add_paragraph(linha_formatada, style='Heading 2')
        elif linha_formatada.startswith("-"):
            doc.add_paragraph(linha_formatada, style='List Bullet')
        else:
            doc.add_paragraph(f"- {linha_formatada}")

    doc.save(file_path)
    messagebox.showinfo("Sucesso", f"Currículo salvo em: {file_path}")

def salvar_como_pdf(texto, imagem_path=None):
    file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Files", "*.pdf")])
    if not file_path:
        return

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    linhas = [corrigir_caracteres(l.strip()) for l in texto.split("\n")]
    if len(linhas) < 2:
        messagebox.showerror("Erro", "Texto inválido.")
        return

    nome, _ = converter_markdown(linhas[0])
    contato = linhas[1]

    if imagem_path:
        try:
            print("Imagem selecionada:", imagem_path)

            pdf.image(imagem_path, x=10, y=10, w=30, h=30)
        except Exception as e:
            print("Erro ao inserir imagem no PDF:", e)

    pdf.set_xy(50, 10)
    pdf.set_font("Arial", "B", 22)
    pdf.cell(0, 10, nome, ln=True)

    pdf.set_xy(50, 20)
    pdf.set_font("Arial", "", 12)
    pdf.cell(0, 10, contato, ln=True)

    pdf.set_y(40)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(8)

    for linha in linhas[2:]:
        if not linha:
            pdf.ln(3)
            continue
        if set(linha) == {"-"}:
            continue
        texto_convertido, is_negrito = converter_markdown(linha)
        if is_negrito:
            pdf.set_font("Arial", "B", 14)
            pdf.multi_cell(0, 8, texto_convertido)
            pdf.ln(4)
        else:
            pdf.set_font("Arial", "", 12)
            pdf.multi_cell(0, 8, texto_convertido)
            pdf.ln(2)

    pdf.output(file_path)
    messagebox.showinfo("Sucesso", f"Currículo salvo em: {file_path}")