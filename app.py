import google.generativeai as genai
from fpdf import FPDF
import tkinter as tk
from tkinter import messagebox, Toplevel, scrolledtext, filedialog
import re
from docx import Document
from docx.shared import Pt


API_KEY = "AIzaSyA90mSpQCPKjt8d-mRRD4gP94GsxVqf9KE"
genai.configure(api_key=API_KEY)

GEMINI_MODEL = "gemini-1.5-flash-latest"

def add_placeholder_to_entry(entry, placeholder_text):
    def on_focus_in(_event):
        if entry.get() == placeholder_text:
            entry.delete(0, tk.END)
            entry.config(fg="black")

    def on_focus_out(_event):
        if entry.get().strip() == "":
            entry.delete(0, tk.END)
            entry.insert(0, placeholder_text)
            entry.config(fg="gray")

    entry.insert(0, placeholder_text)
    entry.config(fg="gray")

    entry.bind("<FocusIn>", on_focus_in)
    entry.bind("<FocusOut>", on_focus_out)

def add_placeholder_to_text(text_widget, placeholder_text):
    def on_focus_in(_event):
        current = text_widget.get("1.0", tk.END).strip()
        if current == placeholder_text:
            text_widget.delete("1.0", tk.END)
            text_widget.config(fg="black")

    def on_focus_out(_event):
        current = text_widget.get("1.0", tk.END).strip()
        if current == "":
            text_widget.delete("1.0", tk.END)
            text_widget.insert("1.0", placeholder_text)
            text_widget.config(fg="gray")

    text_widget.insert("1.0", placeholder_text)
    text_widget.config(fg="gray")

    text_widget.bind("<FocusIn>", on_focus_in)
    text_widget.bind("<FocusOut>", on_focus_out)

def gerar_curriculo_gemini(dados):
    prompt = f"""
    Crie um currículo profissional, bem formatado, usando as informações abaixo:
    Nome: {dados['nome']}
    Cargo Desejado: {dados['cargo']}
    Contato: {dados['contato']}

    Experiência:
    {dados['experiencia']}

    Educação:
    {dados['educacao']}

    Habilidades:
    {dados['habilidades']}

    Projetos:
    {dados['projetos']}

    Idiomas:
    {dados['idiomas']}

    Retorne apenas o currículo, sem mensagens extras, de forma organizada.
    """

    try:
        model = genai.GenerativeModel(model_name=GEMINI_MODEL)
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        print("Erro ao chamar Gemini:", e)
        return "Erro ao gerar currículo com a API do Gemini."

import re

def converter_markdown_pdf(linha):
    if re.match(r"^\*\*(.*?)\*\*$", linha):
        texto_sem_asteriscos = re.sub(r"^\*\*(.*?)\*\*$", r"\1", linha)
        return texto_sem_asteriscos, True
    else:
        texto_sem_asteriscos = re.sub(r"\*\*(.*?)\*\*", r"\1", linha)
        texto_sem_asteriscos = texto_sem_asteriscos.replace("*", "")
        return texto_sem_asteriscos, False

def converter_markdown_docx(linha):
    if re.match(r"^\*\*(.*?)\*\*$", linha):
        texto_sem_asteriscos = re.sub(r"^\*\*(.*?)\*\*$", r"\1", linha)
        return texto_sem_asteriscos, True
    else:
        texto_sem_asteriscos = re.sub(r"\*\*(.*?)\*\*", r"\1", linha)
        texto_sem_asteriscos = texto_sem_asteriscos.replace("*", "")
        return texto_sem_asteriscos, False

def corrigir_caracteres(texto):
    texto = texto.replace("\u2013", "-")
    texto = texto.replace("\u2014", "--")
    texto = texto.replace("•", "-")
    return texto
def salvar_como_docx(texto):
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Documentos do Word", "*.docx")]
    )
    if not file_path:
        return

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    linhas = texto.split("\n")
    linhas = [corrigir_caracteres(l.strip()) for l in linhas]

    if len(linhas) < 2:
        messagebox.showerror("Erro", "Texto inválido.")
        return

    nome, _ = converter_markdown_docx(linhas[0])
    doc.add_heading(nome, level=0)

    contato, _ = converter_markdown_docx(linhas[1])
    doc.add_paragraph(contato)
    doc.add_paragraph("-" * 30)

    for linha in linhas[2:]:
        if not linha.strip():
            doc.add_paragraph()
            continue

        linha_formatada, is_negrito = converter_markdown_docx(linha)
        
        if is_negrito:
            doc.add_paragraph(linha_formatada, style='Heading 2')
        elif linha_formatada.startswith("-") or linha_formatada.startswith("•"):
            doc.add_paragraph(linha_formatada, style='List Bullet')
        else:
            doc.add_paragraph(f"- {linha_formatada}")

    doc.save(file_path)
    messagebox.showinfo("Sucesso", f"Currículo salvo em: {file_path}")
def salvar_como_pdf(texto):
    file_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF Files", "*.pdf")]
    )
    if not file_path:
        return

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    linhas = texto.split("\n")
    linhas = [corrigir_caracteres(l.strip()) for l in linhas]

    if len(linhas) < 2:
        messagebox.showerror("Erro", "Texto inválido.")
        return

    nome, _ = converter_markdown_pdf(linhas[0])
    contato = linhas[1]

    pdf.set_font("Arial", "B", 22)
    pdf.cell(0, 12, nome, ln=True)
    pdf.ln(3)

    pdf.set_font("Arial", "", 12)
    pdf.multi_cell(0, 8, contato)
    pdf.ln(5)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)

    for linha in linhas[2:]:
        if not linha:
            pdf.ln(3)
            continue
        if set(linha) == {"-"}:
            continue

        texto_convertido, is_negrito = converter_markdown_pdf(linha)
        if is_negrito:
            pdf.ln(2)
            pdf.set_font("Arial", "B", 14)
            pdf.multi_cell(0, 8, texto_convertido)
            pdf.ln(4)
        else:
            pdf.set_font("Arial", "", 12)
            pdf.multi_cell(0, 8, texto_convertido)
            pdf.ln(2)

    pdf.output(file_path)
    messagebox.showinfo("Sucesso", f"Currículo salvo em: {file_path}")

def mostrar_previa(texto):
    preview_window = Toplevel(root)
    preview_window.title("Prévia do Currículo")

    label = tk.Label(preview_window, text="Edite seu Currículo Antes de Salvar", font=("Arial", 14, "bold"))
    label.pack(pady=5)

    preview_text = scrolledtext.ScrolledText(preview_window, height=20, width=80, wrap="word", font=("Arial", 12))
    preview_text.pack(padx=10, pady=10)
    preview_text.insert("1.0", texto)

    def salvar_edicao_pdf():
        texto_editado = preview_text.get("1.0", tk.END).strip()
        salvar_como_pdf(texto_editado)
        preview_window.destroy()

    def salvar_edicao_docx():
        texto_editado = preview_text.get("1.0", tk.END).strip()
        salvar_como_docx(texto_editado)
        preview_window.destroy()

    btn_pdf = tk.Button(preview_window, text="Salvar como PDF", font=("Arial", 12), command=salvar_edicao_pdf)
    btn_pdf.pack(pady=5)

    btn_docx = tk.Button(preview_window, text="Salvar como DOCX", font=("Arial", 12), command=salvar_edicao_docx)
    btn_docx.pack(pady=5)


def gerar_e_prever():
    nome = corrigir_caracteres(nome_entry.get().strip())
    cargo = corrigir_caracteres(cargo_entry.get().strip())
    contato = corrigir_caracteres(contato_entry.get().strip())

    exp = experiencia_text.get("1.0", tk.END).strip()
    edu = educacao_text.get("1.0", tk.END).strip()
    hab = habilidades_text.get("1.0", tk.END).strip()
    proj = projetos_text.get("1.0", tk.END).strip()
    idi = idiomas_text.get("1.0", tk.END).strip()

    if not nome or not cargo:
        messagebox.showwarning("Atenção", "Nome e Cargo são obrigatórios!")
        return

    dados = {
        "nome": nome,
        "cargo": cargo,
        "contato": contato,
        "experiencia": exp,
        "educacao": edu,
        "habilidades": hab,
        "projetos": proj,
        "idiomas": idi
    }

    texto_curriculo = gerar_curriculo_gemini(dados)
    if "Erro ao gerar currículo" in texto_curriculo:
        messagebox.showerror("Erro", texto_curriculo)
        return

    mostrar_previa(texto_curriculo)

def preencher_campos_teste():
    if nome_entry.get().startswith("Ex:"):
        nome_entry.delete(0, tk.END)
    nome_entry.config(fg="black")
    nome_entry.insert(0, "José da Silva")

    if cargo_entry.get().startswith("Ex:"):
        cargo_entry.delete(0, tk.END)
    cargo_entry.config(fg="black")
    cargo_entry.insert(0, "Auxiliar Administrativo")

    if contato_entry.get().startswith("Ex:"):
        contato_entry.delete(0, tk.END)
    contato_entry.config(fg="black")
    contato_entry.insert(0, "Av. Afonso Pena, 333, Centro, Diadema\nTelefone: 11 98889-1344\nE-mail: josesilva@aol.com")

    experiencia_text.delete("1.0", tk.END)
    experiencia_text.config(fg="black")
    experiencia_text.insert("1.0", "2019 - atual - Araraquara Transportes\nCargo: Auxiliar Administrativo")

    educacao_text.delete("1.0", tk.END)
    educacao_text.config(fg="black")
    educacao_text.insert("1.0", "2003 - Ensino Médio\nEscola Estadual Antônio Ribeiro")

    habilidades_text.delete("1.0", tk.END)
    habilidades_text.config(fg="black")
    habilidades_text.insert("1.0", "Microsoft Office\nGestão de Tempo\nComunicação")

    projetos_text.delete("1.0", tk.END)
    projetos_text.config(fg="black")
    projetos_text.insert("1.0", "Sistema de Gestão Interna\nRelatórios Gerenciais")

    idiomas_text.delete("1.0", tk.END)
    idiomas_text.config(fg="black")
    idiomas_text.insert("1.0", "Português (Nativo)\nInglês (Intermediário)")

root = tk.Tk()
root.title("Gerador de Currículos com Gemini")

main_frame = tk.Frame(root, padx=10, pady=10)
main_frame.grid(row=0, column=0, sticky="nsew")
root.grid_rowconfigure(0, weight=1)
root.grid_columnconfigure(0, weight=1)

def criar_campo(master, label_text, placeholder, is_text=False, row=0):
    lbl = tk.Label(master, text=label_text, font=("Arial", 12, "bold"))
    lbl.grid(row=row, column=0, sticky="w", pady=5)

    if is_text:
        entry = tk.Text(master, height=5, width=40, font=("Arial", 12))
        entry.grid(row=row, column=1, sticky="ew", pady=5, padx=5)
        add_placeholder_to_text(entry, placeholder)
    else:
        entry = tk.Entry(master, width=40, font=("Arial", 12))
        entry.grid(row=row, column=1, sticky="ew", pady=5, padx=5)
        add_placeholder_to_entry(entry, placeholder)

    return entry

nome_entry = criar_campo(main_frame, "Nome:", "Ex: João da Silva", False, 0)
cargo_entry = criar_campo(main_frame, "Cargo Desejado:", "Ex: Desenvolvedor Python", False, 1)
contato_entry = criar_campo(main_frame, "Contato:", "Ex: Telefone, e-mail, endereço", False, 2)
experiencia_text = criar_campo(main_frame, "Experiência:", "Ex: Empresa X (2015-2019)...", True, 3)
educacao_text = criar_campo(main_frame, "Educação:", "Ex: Graduação em ...", True, 4)
habilidades_text = criar_campo(main_frame, "Habilidades:", "Ex: Python, SQL, Git...", True, 5)
projetos_text = criar_campo(main_frame, "Projetos (Opcional):", "Ex: Sistema de Gestão...", True, 6)
idiomas_text = criar_campo(main_frame, "Idiomas (Opcional):", "Ex: Português, Inglês...", True, 7)

btn_frame = tk.Frame(root, pady=10)
btn_frame.grid(row=1, column=0, sticky="ew")

gerar_btn = tk.Button(btn_frame, text="Gerar Currículo", font=("Arial", 12), command=gerar_e_prever)
gerar_btn.pack(side="left", padx=5)

teste_btn = tk.Button(btn_frame, text="Preencher Teste", font=("Arial", 12), command=preencher_campos_teste)
teste_btn.pack(side="left", padx=5)

for i in range(2):
    main_frame.grid_columnconfigure(i, weight=1)

root.mainloop()